from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from PIL.ExifTags import TAGS
import os
import tempfile
import shutil
from datetime import datetime
import uuid
import re

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max total

# Extensiones válidas de imagen
VALID_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}

def allowed_file(filename):
    return os.path.splitext(filename.lower())[1] in VALID_EXTENSIONS

def extract_image_metadata(filepath):
    """
    Extrae metadata de una imagen (EXIF, nombre del archivo, fecha de modificación)
    Retorna un diccionario con la información disponible
    """
    metadata = {
        'sender': None,
        'datetime': None,
        'filename': os.path.basename(filepath),
        'file_mtime': None
    }
    
    try:
        # Obtener fecha de modificación del archivo
        file_stat = os.stat(filepath)
        metadata['file_mtime'] = datetime.fromtimestamp(file_stat.st_mtime)
        
        # Intentar extraer información del nombre del archivo
        # WhatsApp suele guardar imágenes como: IMG-20231225-WA0001.jpg
        # O con formato de timestamp
        filename = os.path.basename(filepath)
        
        # Buscar patrones de WhatsApp en el nombre
        # Patrón: IMG-YYYYMMDD-WA####
        wsp_pattern = r'IMG-(\d{8})-WA\d+'
        match = re.search(wsp_pattern, filename)
        if match:
            date_str = match.group(1)
            try:
                metadata['datetime'] = datetime.strptime(date_str, '%Y%m%d')
            except:
                pass
        
        # Intentar leer EXIF data
        with Image.open(filepath) as img:
            exif_data = img._getexif()
            
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    
                    # Buscar fecha/hora original
                    if tag == 'DateTimeOriginal' or tag == 'DateTime':
                        try:
                            metadata['datetime'] = datetime.strptime(value, '%Y:%m:%d %H:%M:%S')
                        except:
                            pass
                    
                    # Buscar información del autor/creador
                    elif tag == 'Artist' or tag == 'Author':
                        metadata['sender'] = value
                    
                    # XPAuthor (Windows)
                    elif tag == 'XPAuthor':
                        try:
                            metadata['sender'] = value.decode('utf-16le').rstrip('\x00')
                        except:
                            pass
                    
                    # UserComment puede contener información adicional
                    elif tag == 'UserComment':
                        try:
                            if isinstance(value, bytes):
                                comment = value.decode('utf-8', errors='ignore')
                                metadata['sender'] = comment
                        except:
                            pass
        
        # Si no encontramos fecha en EXIF, usar la fecha de modificación
        if not metadata['datetime']:
            metadata['datetime'] = metadata['file_mtime']
            
    except Exception as e:
        print(f"Error extrayendo metadata de {filepath}: {e}")
    
    return metadata

def sort_images_by_metadata(file_list):
    """
    Ordena imágenes solo por fecha/hora
    file_list: lista de tuplas (filename, filepath)
    Retorna: tupla (lista ordenada de filepath, lista de metadata completa)
    """
    images_with_metadata = []
    
    for filename, filepath in file_list:
        metadata = extract_image_metadata(filepath)
        images_with_metadata.append({
            'filepath': filepath,
            'filename': filename,
            'sender': metadata['sender'] or 'Unknown',
            'datetime': metadata['datetime'] or datetime(1970, 1, 1),
            'metadata': metadata  # Mantener metadata completa
        })
    
    # Ordenar solo por datetime (fecha/hora de envío)
    sorted_images = sorted(images_with_metadata, key=lambda x: x['datetime'])
    
    return [img['filepath'] for img in sorted_images], images_with_metadata

def images_to_word(image_paths, output_file, mode='standard', images_metadata=None):
    """
    Convierte una lista de imágenes a un documento Word
    
    Args:
        image_paths: Lista de rutas de archivos de imagen
        output_file: Ruta del archivo de salida .docx
        mode: 'standard' (1 por página) o 'receipts' (grid 2x2)
        images_metadata: Lista opcional de diccionarios con metadata para cada imagen
    """
    document = Document()
    
    # Crear un mapa de filepath -> metadata para búsqueda rápida
    metadata_map = {}
    if images_metadata:
        for img_data in images_metadata:
            metadata_map[img_data['filepath']] = img_data
    
    # Configurar márgenes según el modo
    section = document.sections[0]
    
    if mode == 'standard':
        # Modo estándar con márgenes normales
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
    else:
        # Modo recibos: SIN MÁRGENES (100% de la hoja)
        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)

    # Dimensiones de página
    page_width = section.page_width
    page_height = section.page_height
    
    # Calcular ancho y alto disponibles
    available_width = page_width - section.left_margin - section.right_margin
    available_height = page_height - section.top_margin - section.bottom_margin

    processed = 0
    errors = []

    if mode == 'standard':
        # MODO ESTÁNDAR: Priorizar visibilidad completa (generalmente 1 por página si son grandes)
        for i, filepath in enumerate(image_paths):
            try:
                # Agregar salto de página antes (excepto en la primera imagen)
                if i > 0:
                    document.add_page_break()
                
                # Agregar fecha/hora si hay metadata disponible
                if filepath in metadata_map:
                    img_metadata = metadata_map[filepath]
                    if img_metadata.get('datetime'):
                        # Agregar párrafo con fecha y hora
                        date_paragraph = document.add_paragraph()
                        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = date_paragraph.add_run(
                            img_metadata['datetime'].strftime('📅 %d/%m/%Y  🕐 %H:%M:%S')
                        )
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(102, 126, 234)  # Color morado
                        
                        # Espacio pequeño entre fecha e imagen
                        date_paragraph.paragraph_format.space_after = Pt(6)
                
                # Agregar la imagen
                with Image.open(filepath) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_width / img_height
                    
                    target_width = available_width
                    target_height = int(target_width / aspect_ratio)
                    
                    if target_height > available_height:
                        target_height = available_height
                        target_width = int(target_height * aspect_ratio)
                    
                    document.add_picture(filepath, width=target_width, height=target_height)
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    processed += 1
            except Exception as e:
                errors.append(f"{os.path.basename(filepath)}: {str(e)}")

    else:
        # MODO RECIBOS: Grid 2 columnas x 2 filas (4 imágenes por hoja)
        # Crear tabla
        table = document.add_table(rows=0, cols=2)
        table.autofit = False 
        
        # Calcular dimensiones de celda para 2 columnas y 2 filas por página
        col_width = available_width / 2
        row_height = available_height / 2
        
        # Sin márgenes internos para ocupar toda la superficie
        cell_margin = Mm(0)
        
        # Dimensiones máximas de imagen dentro de la celda (toda la celda)
        max_img_width = col_width
        max_img_height = row_height

        current_row = None

        for i, filepath in enumerate(image_paths):
            try:
                # Determinar columna (0 o 1)
                col_idx = i % 2
                
                # Si es columna 0, crear nueva fila
                if col_idx == 0:
                    current_row = table.add_row()
                    current_row.height = int(row_height)
                    current_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                cell = current_row.cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.width = int(col_width)
                
                # Limpiar párrafo existente en la celda
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Agregar fecha/hora si hay metadata disponible
                if filepath in metadata_map:
                    img_metadata = metadata_map[filepath]
                    if img_metadata.get('datetime'):
                        # Agregar texto con fecha y hora
                        run = paragraph.add_run(
                            img_metadata['datetime'].strftime('📅 %d/%m/%Y %H:%M\n')
                        )
                        run.font.size = Pt(8)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(102, 126, 234)
                
                # Procesar imagen
                with Image.open(filepath) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_width / img_height
                    
                    # Calcular tamaño para llenar completamente la celda
                    # Reservar espacio para el texto de fecha si existe
                    available_cell_height = max_img_height
                    if filepath in metadata_map and metadata_map[filepath].get('datetime'):
                        available_cell_height = max_img_height - Mm(8)  # Reservar 8mm para la fecha
                    
                    target_width = max_img_width
                    target_height = int(target_width / aspect_ratio)
                    
                    # Si es muy alta, ajustar por alto
                    if target_height > available_cell_height:
                        target_height = available_cell_height
                        target_width = int(target_height * aspect_ratio)
                    
                    run = paragraph.add_run()
                    run.add_picture(filepath, width=int(target_width), height=int(target_height))
                    
                processed += 1

            except Exception as e:
                errors.append(f"{os.path.basename(filepath)}: {str(e)}")

    document.save(output_file)
    return processed, errors

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'images' not in request.files:
        return jsonify({'error': 'No se encontraron imágenes'}), 400
    
    files = request.files.getlist('images')
    mode = request.form.get('mode', 'standard')
    sort_by = request.form.get('sort_by', 'name')  # 'name' o 'metadata'
    
    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'No se seleccionaron archivos'}), 400
    
    # Crear directorio temporal
    temp_dir = tempfile.mkdtemp()
    image_paths = []
    
    try:
        # Guardar archivos temporalmente
        saved_files = []
        for file in files:
            if file and file.filename and allowed_file(file.filename):
                # Mantener nombre original para ordenar
                filename = file.filename
                filepath = os.path.join(temp_dir, filename)
                file.save(filepath)
                saved_files.append((filename, filepath))
        
        if not saved_files:
            return jsonify({'error': 'No se encontraron imágenes válidas'}), 400
        
        # Ordenar según el método seleccionado
        metadata_list = None
        if sort_by == 'metadata':
            # Ordenar por metadata (fecha/hora)
            image_paths, metadata_list = sort_images_by_metadata(saved_files)
        else:
            # Ordenar por nombre de archivo (comportamiento original)
            saved_files.sort(key=lambda x: x[0])
            image_paths = [f[1] for f in saved_files]
        
        # Generar documento Word
        output_filename = f"documento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        processed, errors = images_to_word(image_paths, output_path, mode, metadata_list)
        
        if processed == 0:
            return jsonify({'error': 'No se pudo procesar ninguna imagen'}), 400
        
        # Enviar archivo
        response = send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Limpiar después de enviar (usando callback)
        @response.call_on_close
        def cleanup():
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass
        
        return response
        
    except Exception as e:
        # Limpiar en caso de error
        shutil.rmtree(temp_dir, ignore_errors=True)
        return jsonify({'error': f'Error al procesar: {str(e)}'}), 500

@app.route('/analyze_metadata', methods=['POST'])
def analyze_metadata():
    """Analiza metadata de las imágenes sin convertirlas"""
    if 'images' not in request.files:
        return jsonify({'error': 'No se encontraron imágenes'}), 400
    
    files = request.files.getlist('images')
    
    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'No se seleccionaron archivos'}), 400
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        saved_files = []
        for file in files:
            if file and file.filename and allowed_file(file.filename):
                filename = file.filename
                filepath = os.path.join(temp_dir, filename)
                file.save(filepath)
                saved_files.append((filename, filepath))
        
        if not saved_files:
            return jsonify({'error': 'No se encontraron imágenes válidas'}), 400
        
        # Extraer metadata de todas las imágenes
        metadata_results = []
        for filename, filepath in saved_files:
            metadata = extract_image_metadata(filepath)
            metadata_results.append({
                'filename': filename,
                'sender': metadata['sender'] or 'Desconocido',
                'datetime': metadata['datetime'].strftime('%Y-%m-%d %H:%M:%S') if metadata['datetime'] else 'No disponible',
                'file_mtime': metadata['file_mtime'].strftime('%Y-%m-%d %H:%M:%S') if metadata['file_mtime'] else 'No disponible'
            })
        
        # Limpiar archivos temporales
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        return jsonify({
            'success': True,
            'total_images': len(metadata_results),
            'metadata': metadata_results
        })
        
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return jsonify({'error': f'Error al analizar: {str(e)}'}), 500

if __name__ == '__main__':
    # Crear carpeta templates si no existe
    os.makedirs('templates', exist_ok=True)
    print("🚀 Servidor iniciado en http://localhost:5001")
    app.run(debug=True, host='0.0.0.0', port=5001)
