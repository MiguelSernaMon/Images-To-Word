import os
from docx import Document
from docx.shared import Mm
from PIL import Image
import sys

def images_to_word(image_folder, output_file):
    document = Document()
    
    # Set narrow margins (e.g., 10mm) to maximize space
    section = document.sections[0]
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    # Page dimensions (A4 by default in python-docx)
    page_width = section.page_width
    page_height = section.page_height
    
    # Calculate available width and height
    available_width = page_width - section.left_margin - section.right_margin
    available_height = page_height - section.top_margin - section.bottom_margin

    valid_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff')
    
    # Get list of image files
    try:
        files = [f for f in os.listdir(image_folder) if f.lower().endswith(valid_extensions)]
        files.sort() # Sort alphabetically
    except FileNotFoundError:
        print(f"Error: The folder '{image_folder}' was not found.")
        return

    if not files:
        print(f"No images found in '{image_folder}'.")
        return

    print(f"Found {len(files)} images. Processing...")

    for i, filename in enumerate(files):
        filepath = os.path.join(image_folder, filename)
        
        try:
            # Open image to get dimensions
            with Image.open(filepath) as img:
                img_width, img_height = img.size
                
                # Calculate aspect ratio
                aspect_ratio = img_width / img_height
                
                # Determine target dimensions to fit within available space
                # 1. Try fitting by width
                target_width = available_width
                target_height = int(target_width / aspect_ratio)
                
                # 2. If height exceeds available height, fit by height
                if target_height > available_height:
                    target_height = available_height
                    target_width = int(target_height * aspect_ratio)
                
                # Add image to document
                document.add_picture(filepath, width=target_width, height=target_height)
                
                # Get the paragraph containing the image (it's the last one added)
                last_paragraph = document.paragraphs[-1]
                
                # Center the image (optional but looks better)
                last_paragraph.alignment = 1  # 1 is CENTER

                # If it's not the first image, add a page break before this paragraph
                if i > 0:
                    last_paragraph.paragraph_format.page_break_before = True
                
                print(f"Added {filename}")

                    
        except Exception as e:
            print(f"Failed to process {filename}: {e}")

    try:
        document.save(output_file)
        print(f"Successfully created '{output_file}'")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    # Default to current directory if no arguments provided
    folder = "."
    if len(sys.argv) > 1:
        folder = sys.argv[1]
    
    output = "output.docx"
    if len(sys.argv) > 2:
        output = sys.argv[2]
        
    images_to_word(folder, output)
